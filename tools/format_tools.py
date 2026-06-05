"""
Copiador de formato entre documentos Word.

Toma un documento EJEMPLO (la guía de estilo) y aplica sus formatos a un
documento DESTINO ya escrito, identificando qué párrafo es título, subtítulo,
encabezado o cuerpo (modo híbrido: estilos de Word + heurística) y asignándole
el formato correspondiente del ejemplo.

Estrategia:
  1. Copiar las DEFINICIONES de estilo del ejemplo al destino (docDefaults +
     w:style de Título, Subtítulo, Encabezado 1-3, Normal, lista y tablas).
  2. Copiar los márgenes y la configuración de página.
  3. Reclasificar cada párrafo del destino y asignarle el styleId correcto,
     limpiando el formato directo que tape el estilo.
  4. Aplicar el estilo de tabla del ejemplo a las tablas del destino.
"""

import os
import re
import copy
from collections import Counter

from docx import Document
from docx.oxml.ns import qn
from sandbox import resolve_path

# ─────────────────────────────────────────────────────────────────────
# Niveles canónicos y mapeos
# ─────────────────────────────────────────────────────────────────────

# styleId (independiente del idioma) por nivel canónico
ESTILO_ID_POR_NIVEL = {
    "TITULO": "Title",
    "SUBTITULO": "Subtitle",
    "H1": "Heading1",
    "H2": "Heading2",
    "H3": "Heading3",
    "BODY": "Normal",
}

# styleId del párrafo → nivel canónico (para el modo "por estilo")
NIVEL_POR_STYLEID = {
    "Title": "TITULO",
    "Subtitle": "SUBTITULO",
    "Heading1": "H1",
    "Heading2": "H2",
    "Heading3": "H3",
    "Heading4": "H3",
    "Heading5": "H3",
    "Heading6": "H3",
}

# Nombre de estilo localizado (es/en) → nivel, como respaldo
NOMBRE_A_NIVEL = {
    "title": "TITULO", "título": "TITULO", "titulo": "TITULO",
    "subtitle": "SUBTITULO", "subtítulo": "SUBTITULO", "subtitulo": "SUBTITULO",
    "heading 1": "H1", "título 1": "H1", "titulo 1": "H1",
    "heading 2": "H2", "título 2": "H2", "titulo 2": "H2",
    "heading 3": "H3", "título 3": "H3", "titulo 3": "H3",
}

# styleIds de párrafo a copiar siempre desde el ejemplo
STYLE_IDS_BASE = ["Normal", "Title", "Subtitle", "Heading1", "Heading2",
                  "Heading3", "Heading4", "ListParagraph"]


# ─────────────────────────────────────────────────────────────────────
# Helpers de XML
# ─────────────────────────────────────────────────────────────────────

def _style_id_de_parrafo(par):
    """Devuelve el styleId aplicado directamente al párrafo, o None."""
    pPr = par._p.find(qn("w:pPr"))
    if pPr is not None:
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            return pStyle.get(qn("w:val"))
    return None


def _set_style_id_parrafo(par, style_id):
    """Asigna un styleId al párrafo de forma independiente del idioma."""
    pPr = par._p.get_or_add_pPr()
    for e in pPr.findall(qn("w:pStyle")):
        pPr.remove(e)
    pStyle = pPr.makeelement(qn("w:pStyle"), {qn("w:val"): style_id})
    pPr.insert(0, pStyle)


def _set_table_style_id(tabla, style_id):
    """Asigna un estilo de tabla por styleId."""
    tblPr = tabla._tbl.tblPr
    if tblPr is None:
        return
    for e in tblPr.findall(qn("w:tblStyle")):
        tblPr.remove(e)
    el = tblPr.makeelement(qn("w:tblStyle"), {qn("w:val"): style_id})
    tblPr.insert(0, el)


def _limpiar_run(run, conservar_enfasis=False):
    """
    Elimina el formato DIRECTO de un run para que el estilo del párrafo gobierne.
    Si conservar_enfasis=True, mantiene negrita/cursiva (útil para el cuerpo).
    """
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return
    quitar = ["w:rFonts", "w:sz", "w:szCs", "w:color"]
    if not conservar_enfasis:
        quitar += ["w:b", "w:bCs", "w:i", "w:iCs"]
    for tag in quitar:
        for e in rPr.findall(qn(tag)):
            rPr.remove(e)


# ─────────────────────────────────────────────────────────────────────
# Copia de definiciones de estilo y página
# ─────────────────────────────────────────────────────────────────────

def _estilos_tabla_usados(doc):
    """styleIds de los estilos de tabla usados por las tablas del documento."""
    ids = set()
    for t in doc.tables:
        try:
            if t.style is not None and t.style.style_id:
                ids.add(t.style.style_id)
        except Exception:
            pass
    return ids


def _copiar_definiciones_estilo(doc_ejemplo, doc_destino):
    """Copia docDefaults y las definiciones w:style relevantes del ejemplo."""
    src = doc_ejemplo.styles.element
    dst = doc_destino.styles.element

    # 1. docDefaults (fuente y párrafo por defecto del documento)
    src_defaults = src.find(qn("w:docDefaults"))
    if src_defaults is not None:
        old = dst.find(qn("w:docDefaults"))
        if old is not None:
            dst.remove(old)
        dst.insert(0, copy.deepcopy(src_defaults))

    # 2. Índice de estilos del ejemplo por styleId
    src_por_id = {}
    for st in src.findall(qn("w:style")):
        sid = st.get(qn("w:styleId"))
        if sid:
            src_por_id[sid] = st

    ids_a_copiar = set(STYLE_IDS_BASE) | _estilos_tabla_usados(doc_ejemplo)

    copiados = []
    for sid in ids_a_copiar:
        st = src_por_id.get(sid)
        if st is None:
            continue
        # Quitar la definición previa con el mismo id en el destino
        for ex in dst.findall(qn("w:style")):
            if ex.get(qn("w:styleId")) == sid:
                dst.remove(ex)
        dst.append(copy.deepcopy(st))
        copiados.append(sid)
    return copiados


def _copiar_margenes(doc_ejemplo, doc_destino):
    """Copia márgenes, tamaño y orientación de página del ejemplo al destino."""
    if not doc_ejemplo.sections:
        return
    sec_e = doc_ejemplo.sections[0]
    for sec in doc_destino.sections:
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin",
                     "header_distance", "footer_distance",
                     "page_width", "page_height", "orientation"):
            try:
                valor = getattr(sec_e, attr)
                if valor is not None:
                    setattr(sec, attr, valor)
            except Exception:
                pass


# ─────────────────────────────────────────────────────────────────────
# Clasificación de párrafos (híbrido: estilo + heurística)
# ─────────────────────────────────────────────────────────────────────

def _tamano_par(par, default):
    """Tamaño de fuente (pt) del párrafo: mayor de sus runs o el de su estilo."""
    tam = [r.font.size.pt for r in par.runs if r.font.size is not None]
    if tam:
        return max(tam)
    try:
        if par.style is not None and par.style.font.size is not None:
            return par.style.font.size.pt
    except Exception:
        pass
    return default


def _es_negrita(par):
    """True si algún run del párrafo está en negrita (formato directo)."""
    return any(r.bold for r in par.runs)


def _tamano_cuerpo(doc, default=11):
    """
    Estima el tamaño de fuente (pt) del cuerpo del documento.
    Para no sesgarse con los encabezados, prioriza los párrafos 'de cuerpo'
    (largos o que terminan en punto). Si no hay tamaños explícitos, usa el
    tamaño del estilo Normal o el valor por defecto.
    """
    c = Counter()
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue
        es_cuerpo = len(texto.split()) > 18 or texto.endswith((".", "!", "?"))
        if not es_cuerpo:
            continue
        for r in p.runs:
            if r.font.size is not None and r.text.strip():
                c[round(r.font.size.pt)] += len(r.text)
    if c:
        return c.most_common(1)[0][0]
    # Respaldo: tamaño del estilo Normal
    try:
        if doc.styles["Normal"].font.size is not None:
            return doc.styles["Normal"].font.size.pt
    except Exception:
        pass
    return default


def _clasificar_parrafo(par, body_size):
    """
    Determina el nivel canónico de un párrafo.
    Modo híbrido: primero por estilo de Word, luego por heurística.
    """
    # 1. Por styleId directo
    sid = _style_id_de_parrafo(par)
    if sid in NIVEL_POR_STYLEID:
        return NIVEL_POR_STYLEID[sid]
    if sid == "Normal":
        # Está marcado como Normal: respetar como cuerpo salvo heurística fuerte
        pass

    # 2. Por nombre de estilo localizado
    try:
        nombre = (par.style.name or "").lower() if par.style else ""
    except Exception:
        nombre = ""
    for clave, niv in NOMBRE_A_NIVEL.items():
        if nombre == clave or nombre.startswith(clave):
            return niv

    # 3. Heurística (para texto plano sin estilos)
    texto = par.text.strip()
    if not texto:
        return None
    palabras = len(texto.split())
    size = _tamano_par(par, body_size)
    bold = _es_negrita(par)
    termina_frase = texto.endswith((".", "!", "?"))

    # 3a. Numeración jerárquica: "1." → H1, "1.1" → H2, "1.1.1" → H3
    #     Se cuenta la profundidad SOLO sobre el prefijo numérico (group 1),
    #     para que el punto separador final no infle el nivel.
    m = re.match(r"^(\d+(?:\.\d+)*)[\.\)]?\s+\S", texto)
    if m and palabras <= 16 and not termina_frase:
        profundidad = m.group(1).count(".")
        if profundidad <= 0:
            return "H1"
        elif profundidad == 1:
            return "H2"
        else:
            return "H3"

    # 3b. ¿Parece encabezado? Corto, sin punto final, negrita o fuente mayor
    es_heading = (
        palabras <= 14
        and not termina_frase
        and (bold or size > body_size * 1.08)
    )
    if not es_heading:
        return "BODY"

    if size >= body_size * 1.5:
        return "TITULO"
    if size >= body_size * 1.25:
        return "H1"
    if size >= body_size * 1.1:
        return "H2"
    return "H3"


def _aplicar_nivel(par, nivel):
    """Asigna el estilo del nivel al párrafo y limpia el formato directo."""
    sid = ESTILO_ID_POR_NIVEL.get(nivel, "Normal")
    _set_style_id_parrafo(par, sid)
    conservar_enfasis = (nivel == "BODY")  # en el cuerpo respetar negrita/cursiva inline
    for run in par.runs:
        _limpiar_run(run, conservar_enfasis=conservar_enfasis)


# ─────────────────────────────────────────────────────────────────────
# Localización de archivos
# ─────────────────────────────────────────────────────────────────────

def _buscar_archivo(filename, contextos):
    """Devuelve la primera ruta existente del archivo entre los contextos dados."""
    for ctx in contextos:
        try:
            ruta = resolve_path(filename, ctx)
            if os.path.exists(ruta):
                return ruta
        except Exception:
            continue
    return None


# ─────────────────────────────────────────────────────────────────────
# Función pública
# ─────────────────────────────────────────────────────────────────────

def copiar_formato_word(ejemplo: str, destino: str, salida: str = None,
                        context_salida: str = "outputs") -> str:
    """
    Aplica el formato del documento `ejemplo` al documento `destino`.

    Args:
        ejemplo: nombre del .docx con el formato a copiar (se busca en
                 templates/, luego inputs/ y outputs/).
        destino: nombre del .docx con el contenido a reformatear (se busca en
                 inputs/, luego templates/ y outputs/).
        salida:  nombre del archivo resultante (por defecto
                 'formateado_<destino>.docx' en outputs/).

    Returns:
        Mensaje con el resultado o un texto que empieza por 'Error'.
    """
    ruta_ej = _buscar_archivo(ejemplo, ["templates", "inputs", "outputs"])
    if not ruta_ej:
        return (f"Error: No se encontró el documento de ejemplo '{ejemplo}'. "
                f"Colócalo en la carpeta 'templates/'.")

    ruta_de = _buscar_archivo(destino, ["inputs", "templates", "outputs"])
    if not ruta_de:
        return (f"Error: No se encontró el documento destino '{destino}'. "
                f"Colócalo en la carpeta 'inputs/'.")

    try:
        doc_e = Document(ruta_ej)
        doc_d = Document(ruta_de)
    except Exception as e:
        return f"Error abriendo los documentos Word: {e}"

    try:
        # 1. Copiar definiciones de estilo y página
        _copiar_definiciones_estilo(doc_e, doc_d)
        _copiar_margenes(doc_e, doc_d)

        # 2. Reclasificar y aplicar formato a cada párrafo
        body_size = _tamano_cuerpo(doc_d)
        conteo = Counter()
        for par in doc_d.paragraphs:
            nivel = _clasificar_parrafo(par, body_size)
            if nivel is None:
                continue
            _aplicar_nivel(par, nivel)
            conteo[nivel] += 1

        # 3. Estilo de tabla del ejemplo (si tiene tablas con estilo)
        estilos_tabla = list(_estilos_tabla_usados(doc_e))
        if estilos_tabla:
            for t in doc_d.tables:
                _set_table_style_id(t, estilos_tabla[0])

        # 4. Guardar
        if not salida:
            salida = f"formateado_{os.path.splitext(os.path.basename(destino))[0]}.docx"
        ruta_out = resolve_path(salida, context_salida)
        doc_d.save(ruta_out)

        resumen = ", ".join(f"{niv}:{n}" for niv, n in conteo.items()) or "sin párrafos"
        return (f"Formato copiado de '{ejemplo}' a '{destino}'.\n"
                f"Párrafos reformateados → {resumen}.\n"
                f"Tablas afectadas: {len(doc_d.tables)}.\n"
                f"Guardado en: {ruta_out}")

    except Exception as e:
        return f"Error copiando el formato: {e}"
