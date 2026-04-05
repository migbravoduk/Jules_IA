from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from sandbox import resolve_path
import os
import json


# ─────────────────────────────────────────────
# Helpers APA
# ─────────────────────────────────────────────

def _aplicar_margenes_apa(doc: Document):
    """Configura márgenes APA (2.54 cm = 1 pulgada en todos los lados)."""
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)


def _fuente_parrafo(run, size_pt: int = 12, bold: bool = False, italic: bool = False):
    """Aplica fuente Times New Roman APA a un run."""
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic


def _parrafo_apa(doc: Document, texto: str, nivel: str = "body") -> None:
    """
    Añade un párrafo con formato APA según el nivel:
      body      → Times New Roman 12, justificado, interlineado doble, sangría primera línea
      titulo    → centrado, negrita, 14pt
      subtitulo → negrita, izquierda, 12pt
      cita      → cursiva, centrado, 11pt
    """
    p = doc.add_paragraph()

    if nivel == "titulo":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        _fuente_parrafo(run, size_pt=14, bold=True)
        return

    if nivel == "subtitulo":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(texto)
        _fuente_parrafo(run, size_pt=12, bold=True)
        return

    if nivel == "cita":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'"{texto}"')
        _fuente_parrafo(run, size_pt=11, italic=True)
        return

    # body: justificado, doble espacio, sangría primera línea
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._p.get_or_add_pPr()

    # Interlineado doble
    spacing = p._p.get_or_add_pPr().get_or_add_spacing()
    spacing.set(qn("w:line"), "480")          # 240 = simple, 480 = doble
    spacing.set(qn("w:lineRule"), "auto")

    # Sangría primera línea 0.5"
    ind = pPr.get_or_add_ind()
    ind.set(qn("w:firstLine"), "720")         # 720 twips = 0.5 pulgadas

    run = p.add_run(texto)
    _fuente_parrafo(run, size_pt=12)


def _agregar_referencia(doc: Document, refs: list):
    """Añade la sección 'Referencias' al final del documento en formato APA."""
    # Encabezado centrado
    p_heading = doc.add_paragraph()
    p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_heading.add_run("Referencias")
    _fuente_parrafo(r, size_pt=12, bold=True)

    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Sangría francesa: cuerpo sangrado, primera línea a margen
        pPr = p._p.get_or_add_pPr()
        ind = pPr.get_or_add_ind()
        ind.set(qn("w:left"),      "720")   # 0.5" indent al cuerpo
        ind.set(qn("w:hanging"),   "720")   # primera línea vuelve al margen

        # Interlineado doble
        spacing = pPr.get_or_add_spacing()
        spacing.set(qn("w:line"),     "480")
        spacing.set(qn("w:lineRule"), "auto")

        r = p.add_run(ref)
        _fuente_parrafo(r, size_pt=12)

    # Disclaimer
    p_disc = doc.add_paragraph()
    r_disc = p_disc.add_run(
        "⚠ Nota: Las referencias fueron generadas por IA. Verifique su exactitud antes de publicar."
    )
    _fuente_parrafo(r_disc, size_pt=10, italic=True)
    p_disc.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ─────────────────────────────────────────────
# Funciones públicas
# ─────────────────────────────────────────────

def crear_documento_profesional(titulo: str, contenido: str, filename: str, context: str = "outputs") -> str:
    """Crea un documento Word con formato APA básico desde texto plano."""
    try:
        doc = Document()
        _aplicar_margenes_apa(doc)

        _parrafo_apa(doc, titulo, nivel="titulo")
        doc.add_paragraph()  # línea en blanco

        for linea in contenido.split("\n"):
            linea = linea.strip()
            if not linea:
                continue
            if linea.startswith("- ") or linea.startswith("* "):
                p = doc.add_paragraph(linea[2:], style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif linea.startswith("#"):
                _parrafo_apa(doc, linea.lstrip("#").strip(), nivel="subtitulo")
            else:
                _parrafo_apa(doc, linea, nivel="body")

        filepath = resolve_path(filename, context)
        doc.save(filepath)
        return f"Documento Word APA creado con éxito: {filepath}"
    except Exception as e:
        return f"Error creando documento Word: {e}"


def aplicar_formato_word(filename: str, context: str = "outputs") -> str:
    """Lee un documento existente y aplica alineación justificada a todos los párrafos."""
    try:
        filepath = resolve_path(filename, context)
        if not os.path.exists(filepath):
            return f"Error: No se encontró {filepath}"

        doc = Document(filepath)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        nuevo = os.path.splitext(filename)[0] + "_formateado.docx"
        doc.save(resolve_path(nuevo, context))
        return f"Formato aplicado. Guardado como: {nuevo}"
    except Exception as e:
        return f"Error aplicando formato a Word: {e}"


def crear_word_complejo_desde_json(filename: str, json_str: str, context: str = "outputs") -> str:
    """
    Crea un documento Word con formato APA completo a partir de JSON estructurado.

    Tipos de bloque soportados:
      titulo      → Título principal centrado, 14pt negrita
      subtitulo   → Subtítulo izquierda negrita 12pt
      parrafo     → Párrafo body APA (doble espacio, sangría)
      cita        → Cita destacada en cursiva
      lista       → Lista de viñetas justificada
      tabla       → Tabla con encabezados en negrita
      referencias → Sección Referencias APA con sangría francesa
    """
    try:
        data = json.loads(json_str)
        if not data:
            return "Error: JSON vacío o mal formado."

        doc = Document()
        _aplicar_margenes_apa(doc)

        referencias_pendientes = []   # Se vuelcan al final

        for bloque in data:
            tipo = bloque.get("tipo", "parrafo")

            if tipo == "titulo":
                _parrafo_apa(doc, bloque.get("texto", ""), nivel="titulo")
                doc.add_paragraph()

            elif tipo == "subtitulo":
                doc.add_paragraph()
                _parrafo_apa(doc, bloque.get("texto", ""), nivel="subtitulo")

            elif tipo == "cita":
                doc.add_paragraph()
                _parrafo_apa(doc, bloque.get("texto", ""), nivel="cita")
                doc.add_paragraph()

            elif tipo in ("lista", "indice"):
                items = bloque.get("items", bloque.get("viñetas", []))
                for item in items:
                    p = doc.add_paragraph(item, style="List Bullet")
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        _fuente_parrafo(run, size_pt=12)

            elif tipo == "tabla":
                filas = bloque.get("filas", [])
                titulo_tabla = bloque.get("titulo", "")
                if titulo_tabla:
                    _parrafo_apa(doc, titulo_tabla, nivel="subtitulo")

                if filas and len(filas) > 0:
                    num_cols = len(filas[0])
                    table = doc.add_table(rows=1, cols=num_cols)
                    table.style = "Table Grid"

                    hdr_cells = table.rows[0].cells
                    for i, encabezado in enumerate(filas[0]):
                        hdr_cells[i].text = str(encabezado)
                        for para in hdr_cells[i].paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(12)

                    for fila_datos in filas[1:]:
                        row_cells = table.add_row().cells
                        for i, valor in enumerate(fila_datos):
                            if i < num_cols:
                                row_cells[i].text = str(valor)
                                for para in row_cells[i].paragraphs:
                                    for run in para.runs:
                                        run.font.name = "Times New Roman"
                                        run.font.size = Pt(11)
                doc.add_paragraph()

            elif tipo == "referencias":
                # Acumular para poner al final del documento
                items = bloque.get("items", [])
                referencias_pendientes.extend(items)

            else:  # parrafo por defecto
                _parrafo_apa(doc, bloque.get("texto", ""), nivel="body")

        # Sección Referencias al final (si hay)
        if referencias_pendientes:
            doc.add_page_break()
            _agregar_referencia(doc, referencias_pendientes)

        filepath = resolve_path(filename, context)
        doc.save(filepath)
        return f"Documento Word APA creado con éxito: {filepath}"

    except json.JSONDecodeError as e:
        return f"Error parseando JSON de Word: {e}\nJSON Recibido:\n{json_str[:500]}"
    except Exception as e:
        return f"Error creando documento Word complejo: {e}"
